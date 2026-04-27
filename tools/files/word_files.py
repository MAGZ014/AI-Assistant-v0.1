# tools/files/word_files.py
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import DOCS_FOLDER


def _build_path(filename: str) -> tuple:
    """Asegura extensión .docx y regresa (ruta, filename)."""
    if not filename.endswith(".docx"):
        filename += ".docx"
    return os.path.join(DOCS_FOLDER, filename), filename


# ─────────────────────────────────────────
#  CREAR
# ─────────────────────────────────────────

def create_docx(filename: str, content: str = "", title: str = None) -> str:
    """
    Crea un archivo Word con contenido.
    Soporta secciones con '##', listas con '-' y párrafos normales.
    """
    path, filename = _build_path(filename)
    try:
        doc = Document()

        # ── Estilos base ──
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # ── Título principal ──
        if title:
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif filename:
            heading = doc.add_heading(filename.replace(".docx", "").replace("_", " ").title(), level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ── Fecha de creación ──
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f"Creado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        date_run.font.size = Pt(9)
        date_run.font.color.rgb = RGBColor(128, 128, 128)
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph()  # espacio

        # ── Contenido ──
        _parse_content(doc, content)

        doc.save(path)
        return f"Documento Word '{filename}' creado correctamente."
    except Exception as e:
        return f"Error creando '{filename}': {str(e)}"


def create_docx_with_sections(filename: str, sections: list) -> str:
    """
    Crea un Word con múltiples secciones.
    sections = [{"title": "Introducción", "content": "texto..."}, ...]
    """
    path, filename = _build_path(filename)
    try:
        doc = Document()

        # Título del documento
        doc.add_heading(filename.replace(".docx", "").replace("_", " ").title(), level=1)
        doc.add_paragraph(f"Creado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph()

        for section in sections:
            if section.get("title"):
                doc.add_heading(section["title"], level=2)
            if section.get("content"):
                _parse_content(doc, section["content"])
            doc.add_paragraph()

        doc.save(path)
        return f"Documento '{filename}' creado con {len(sections)} secciones."
    except Exception as e:
        return f"Error creando '{filename}': {str(e)}"


# ─────────────────────────────────────────
#  LEER
# ─────────────────────────────────────────

def read_docx(filename: str, max_chars: int = 3000) -> str:
    """Lee un archivo .docx y regresa su contenido como texto."""
    path, filename = _build_path(filename)

    if not os.path.exists(path):
        return f"No encontré '{filename}'."

    try:
        doc = Document(path)
        lines = []

        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            # Detecta encabezados para darles formato visual en consola
            if para.style.name.startswith("Heading"):
                level = para.style.name.split(" ")[-1]
                prefix = "═" * 30 if level == "1" else "─" * 20
                lines.append(f"\n{prefix}")
                lines.append(para.text.upper() if level == "1" else para.text)
                lines.append(prefix)
            else:
                lines.append(para.text)

        text = "\n".join(lines)

        if len(text) > max_chars:
            text = text[:max_chars] + f"\n\n... [truncado, {len(text) - max_chars} caracteres restantes]"

        return f"Contenido de '{filename}':\n{text}"
    except Exception as e:
        return f"Error leyendo '{filename}': {str(e)}"


# ─────────────────────────────────────────
#  MODIFICAR
# ─────────────────────────────────────────

def append_docx(filename: str, content: str, section_title: str = None) -> str:
    """Agrega contenido al final de un .docx existente."""
    path, filename = _build_path(filename)

    if not os.path.exists(path):
        return f"No encontré '{filename}'. ¿Quieres crearlo primero?"

    try:
        doc = Document(path)

        if section_title:
            doc.add_heading(section_title, level=2)

        _parse_content(doc, content)
        doc.save(path)
        return f"Contenido agregado a '{filename}'."
    except Exception as e:
        return f"Error modificando '{filename}': {str(e)}"


def replace_in_docx(filename: str, old_text: str, new_text: str) -> str:
    """Reemplaza texto dentro de un .docx."""
    path, filename = _build_path(filename)

    if not os.path.exists(path):
        return f"No encontré '{filename}'."

    try:
        doc = Document(path)
        replaced = 0

        for para in doc.paragraphs:
            if old_text in para.text:
                for run in para.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
                        replaced += 1

        doc.save(path)
        return f"{replaced} reemplazo(s) realizados en '{filename}'."
    except Exception as e:
        return f"Error modificando '{filename}': {str(e)}"


# ─────────────────────────────────────────
#  HELPER INTERNO
# ─────────────────────────────────────────

def _parse_content(doc: Document, content: str):
    """
    Parsea el contenido y lo agrega al documento con formato:
    - Líneas que empiezan con ## → Heading 2
    - Líneas que empiezan con ### → Heading 3
    - Líneas que empiezan con - o * → lista con viñeta
    - Resto → párrafo normal
    """
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("1. ") or (len(line) > 2 and line[0].isdigit() and line[1] == "."):
            doc.add_paragraph(line[3:] if line[1] == "." else line[2:], style="List Number")
        else:
            doc.add_paragraph(line)