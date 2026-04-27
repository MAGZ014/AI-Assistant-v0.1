import os
import shutil
from datetime import datetime
from docx import Document
from config import DOCS_FOLDER

# Asegura que la carpeta de documentos exista
os.makedirs(DOCS_FOLDER, exist_ok=True)


# ─────────────────────────────────────────
#  CREAR ARCHIVOS
# ─────────────────────────────────────────

def create_txt(filename: str, content: str = "") -> str:
    """Crea un archivo .txt en la carpeta de documentos."""
    if not filename.endswith(".txt"):
        filename += ".txt"
    path = os.path.join(DOCS_FOLDER, filename)
    try:
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        return f"Archivo '{filename}' creado en {DOCS_FOLDER}."
    except Exception as e:
        return f"Error creando el archivo: {str(e)}"


def create_md(filename: str, content: str = "") -> str:
    """Crea un archivo .md en la carpeta de documentos."""
    if not filename.endswith(".md"):
        filename += ".md"
    path = os.path.join(DOCS_FOLDER, filename)
    try:
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        return f"Archivo '{filename}' creado en {DOCS_FOLDER}."
    except Exception as e:
        return f"Error creando el archivo: {str(e)}"


def create_docx(filename: str, content: str = "") -> str:
    """Crea un archivo Word .docx en la carpeta de documentos."""
    if not filename.endswith(".docx"):
        filename += ".docx"
    path = os.path.join(DOCS_FOLDER, filename)
    try:
        doc = Document()
        # Título = nombre del archivo sin extensión
        doc.add_heading(filename.replace(".docx", ""), level=1)
        # Agrega el contenido como párrafos
        for line in content.split("\n"):
            if line.strip():
                doc.add_paragraph(line)
        doc.save(path)
        return f"Documento Word '{filename}' creado en {DOCS_FOLDER}."
    except Exception as e:
        return f"Error creando el Word: {str(e)}"


# ─────────────────────────────────────────
#  LEER ARCHIVOS
# ─────────────────────────────────────────

def read_file(filename: str) -> str:
    """Lee un archivo .txt, .md o .docx y devuelve su contenido."""
    path = os.path.join(DOCS_FOLDER, filename)

    if not os.path.exists(path):
        return f"No encontré el archivo '{filename}' en {DOCS_FOLDER}."

    try:
        if filename.endswith(".docx"):
            doc = Document(path)
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            return f"Contenido de '{filename}':\n{text}"
        else:
            with open(path, "r", encoding="utf-8") as f:
                text = f.read()
            return f"Contenido de '{filename}':\n{text}"
    except Exception as e:
        return f"Error leyendo el archivo: {str(e)}"


def list_files() -> str:
    """Lista todos los archivos en la carpeta de documentos."""
    try:
        files = os.listdir(DOCS_FOLDER)
        if not files:
            return "La carpeta de documentos está vacía."
        files_list = "\n".join([f"• {f}" for f in files])
        return f"Archivos en {DOCS_FOLDER}:\n{files_list}"
    except Exception as e:
        return f"Error listando archivos: {str(e)}"


# ─────────────────────────────────────────
#  MOVER Y RENOMBRAR
# ─────────────────────────────────────────

def move_file(filename: str, destination: str) -> str:
    """
    Mueve un archivo a otra carpeta.
    destination puede ser ruta absoluta o relativa a DOCS_FOLDER.
    """
    source_path = os.path.join(DOCS_FOLDER, filename)

    if not os.path.exists(source_path):
        return f"No encontré '{filename}' en {DOCS_FOLDER}."

    # Si destination no es absoluta, la trata como subcarpeta dentro de DOCS_FOLDER
    if not os.path.isabs(destination):
        destination = os.path.join(DOCS_FOLDER, destination)

    try:
        os.makedirs(destination, exist_ok=True)
        shutil.move(source_path, os.path.join(destination, filename))
        return f"'{filename}' movido a '{destination}'."
    except Exception as e:
        return f"Error moviendo el archivo: {str(e)}"


def rename_file(filename: str, new_name: str) -> str:
    """Renombra un archivo dentro de la carpeta de documentos."""
    old_path = os.path.join(DOCS_FOLDER, filename)
    new_path = os.path.join(DOCS_FOLDER, new_name)

    if not os.path.exists(old_path):
        return f"No encontré '{filename}' en {DOCS_FOLDER}."

    try:
        os.rename(old_path, new_path)
        return f"'{filename}' renombrado a '{new_name}'."
    except Exception as e:
        return f"Error renombrando el archivo: {str(e)}"


def delete_file(filename: str) -> str:
    """Elimina un archivo de la carpeta de documentos."""
    path = os.path.join(DOCS_FOLDER, filename)

    if not os.path.exists(path):
        return f"No encontré '{filename}'."

    try:
        os.remove(path)
        return f"'{filename}' eliminado."
    except Exception as e:
        return f"Error eliminando el archivo: {str(e)}"